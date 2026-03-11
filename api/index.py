from flask import Flask, request, jsonify, render_template, Response, stream_with_context
import requests
import json
import time
import random
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(BASE_DIR)  # one level up from api/ to root

app = Flask(
    __name__,
    template_folder=ROOT_DIR,
    static_folder=os.path.join(ROOT_DIR, 'static'),
)

# ── Hugging Face Config ────────────────────────────────────────────
HF_API_KEY = os.environ.get("HF_API_KEY", "")
HF_ROUTER_URL = "https://router.huggingface.co/v1/chat/completions"

WORKING_MODELS = [
    "Qwen/Qwen2.5-72B-Instruct",
    "deepseek-ai/DeepSeek-R1-Distill-Llama-70B",
    "deepseek-ai/DeepSeek-V3",
    "meta-llama/Llama-3.1-70B-Instruct",
    "meta-llama/Meta-Llama-3-70B-Instruct",
    "google/gemma-3-27b-it",
    "Qwen/Qwen2.5-7B-Instruct",
    "Qwen/Qwen2.5-Coder-32B-Instruct",
    "Qwen/Qwen2.5-Coder-7B-Instruct",
    "meta-llama/Llama-3.1-8B-Instruct",
    "meta-llama/Llama-3.2-1B-Instruct",
]

MODEL_CATEGORIES = {
    "best_quality": ["Qwen/Qwen2.5-72B-Instruct", "deepseek-ai/DeepSeek-V3", "meta-llama/Llama-3.1-70B-Instruct"],
    "balanced":     ["Qwen/Qwen2.5-7B-Instruct", "meta-llama/Llama-3.1-8B-Instruct"],
    "fast":         ["meta-llama/Llama-3.2-1B-Instruct"],
    "code":         ["Qwen/Qwen2.5-Coder-32B-Instruct", "Qwen/Qwen2.5-Coder-7B-Instruct"],
}

model_stats = {m: {"successes": 0, "failures": 0, "avg_time": 0} for m in WORKING_MODELS}
rate_limit = {}
RATE_LIMIT_SECONDS = 10

SYSTEM_BASE = (
    "You are an expert academic research assistant helping university students "
    "write dissertations. You specialise in African and Zimbabwean research contexts. "
    "Always use formal academic English. Be thorough, structured, and precise."
)

# ── CORS ───────────────────────────────────────────────────────────
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
    response.headers.add('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
    response.headers.add('X-Accel-Buffering', 'no')
    response.headers.add('Cache-Control', 'no-cache')
    return response

# ── Helpers ────────────────────────────────────────────────────────
def check_rate_limit(ip):
    now = time.time()
    if now - rate_limit.get(ip, 0) < RATE_LIMIT_SECONDS:
        return False
    rate_limit[ip] = now
    return True

def get_best_model(task="balanced"):
    candidates = MODEL_CATEGORIES.get(task, MODEL_CATEGORIES["balanced"])
    best_model, best_score = None, -1
    for model in candidates:
        s = model_stats[model]
        total = s["successes"] + s["failures"]
        success_rate = (s["successes"] / total) if total > 0 else 0.5
        time_score   = 1.0 / max(s["avg_time"], 0.1) if s["avg_time"] > 0 else 1.0
        score = success_rate * 0.7 + time_score * 0.3
        if score > best_score:
            best_score, best_model = score, model
    return best_model or random.choice(candidates)

def call_hf_stream(system_msg, user_msg, task="balanced", max_tokens=2000):
    model = get_best_model(task)
    headers = {
        "Authorization": f"Bearer {HF_API_KEY}",
        "Content-Type": "application/json"
    }
    messages = []
    if system_msg:
        messages.append({"role": "system", "content": system_msg})
    messages.append({"role": "user", "content": user_msg})

    payload = {
        "model": model,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "stream": True
    }

    start = time.time()
    try:
        resp = requests.post(HF_ROUTER_URL, headers=headers, json=payload, stream=True, timeout=55)
        elapsed = time.time() - start
        if resp.status_code == 200:
            model_stats[model]["successes"] += 1
            n = model_stats[model]["successes"]
            model_stats[model]["avg_time"] = (model_stats[model]["avg_time"] * (n - 1) + elapsed) / n
            return resp, model
        else:
            model_stats[model]["failures"] += 1
            raise Exception(f"API Error {resp.status_code}: {resp.text[:200]}")
    except Exception as e:
        model_stats[model]["failures"] += 1
        raise

def stream_generator(system_msg, user_msg, task="balanced", max_tokens=2000):
    try:
        hf_response, used_model = call_hf_stream(system_msg, user_msg, task, max_tokens)
        print(f"✅ Using model: {used_model}")

        buffer = ""
        for chunk in hf_response.iter_content(chunk_size=1024, decode_unicode=True):
            if not chunk:
                continue
            buffer += chunk

            while '\n' in buffer:
                newline_pos = buffer.find('\n')
                line = buffer[:newline_pos].strip()
                buffer = buffer[newline_pos + 1:]

                if not line.startswith('data: '):
                    continue

                raw = line[6:]
                if raw == '[DONE]':
                    break

                try:
                    obj = json.loads(raw)
                    if "choices" in obj and obj["choices"]:
                        content = obj["choices"][0].get("delta", {}).get("content")
                        if content:
                            yield f"data: {json.dumps({'content': content})}\n\n"
                except json.JSONDecodeError:
                    pass

        yield "data: [DONE]\n\n"

    except Exception as e:
        print(f"Stream error: {e}")
        yield f"data: {json.dumps({'error': str(e)})}\n\n"

def sse_endpoint(generate_fn):
    if request.method == "OPTIONS":
        return "", 200
    ip = request.remote_addr
    if not check_rate_limit(ip):
        return jsonify({"error": f"Please wait {RATE_LIMIT_SECONDS} seconds between requests"}), 429
    return Response(stream_with_context(generate_fn()), mimetype='text/event-stream')

# ── Routes ─────────────────────────────────────────────────────────
@app.route("/")
def home():
    return render_template("index.html")

# 1. Topic Generator
@app.route("/api/topics", methods=["POST", "OPTIONS"])
def topics():
    def gen():
        body     = request.json or {}
        field    = body.get("field", "").strip()
        region   = body.get("region", "Zimbabwe").strip()
        keywords = body.get("keywords", "").strip()

        if not field:
            yield f"data: {json.dumps({'error': 'Field is required'})}\n\n"
            return

        prompt = f"""Generate 5 dissertation topics for:

Field: {field}
Region: {region}
Keywords: {keywords if keywords else 'general'}

Return ONLY a numbered list (1. to 5.). Each topic should be specific, academic, and researchable."""

        yield from stream_generator(SYSTEM_BASE, prompt, task="balanced", max_tokens=500)

    return sse_endpoint(gen)

# 2. Research Questions
@app.route("/api/questions", methods=["POST", "OPTIONS"])
def questions():
    def gen():
        body  = request.json or {}
        topic = body.get("topic", "").strip()

        if not topic:
            yield f"data: {json.dumps({'error': 'Topic is required'})}\n\n"
            return

        prompt = f"""Generate research questions for this dissertation topic:

Topic: {topic}

Provide:
1. One clear main research question (RQ1)
2. Four supporting sub-questions (SQ1–SQ4)

Format as a numbered list. Each question must be specific, answerable through research, and relevant to the African academic context."""

        yield from stream_generator(SYSTEM_BASE, prompt, task="balanced", max_tokens=500)

    return sse_endpoint(gen)

# 3. Outline Builder
@app.route("/api/outline", methods=["POST", "OPTIONS"])
def outline():
    def gen():
        body   = request.json or {}
        topic  = body.get("topic", "").strip()
        degree = body.get("degree", "Undergraduate").strip()

        if not topic:
            yield f"data: {json.dumps({'error': 'Topic is required'})}\n\n"
            return

        prompt = f"""Create a full dissertation outline for a {degree} dissertation:

Topic: {topic}

Structure it as 5 chapters. For each chapter provide:
- Chapter title
- 4–5 section headings
- Brief one-sentence description per section
- Suggested word count

Use clear academic formatting."""

        yield from stream_generator(SYSTEM_BASE, prompt, task="best_quality", max_tokens=1200)

    return sse_endpoint(gen)

# 4. Section Writer
@app.route("/api/write", methods=["POST", "OPTIONS"])
def write():
    def gen():
        body         = request.json or {}
        topic        = body.get("topic", "").strip()
        section      = body.get("section", "").strip()
        instructions = body.get("instructions", "").strip()

        if not topic or not section:
            yield f"data: {json.dumps({'error': 'Topic and section are required'})}\n\n"
            return

        extra = f"\nAdditional instructions: {instructions}" if instructions else ""
        prompt = f"""Write the '{section}' section for a dissertation on:

Topic: {topic}{extra}

Write in formal academic style with well-structured paragraphs, academic language, and where relevant, cite types of sources a student should reference. Minimum 400 words."""

        yield from stream_generator(SYSTEM_BASE, prompt, task="best_quality", max_tokens=1500)

    return sse_endpoint(gen)

# 5. Methodology Builder
@app.route("/api/methodology", methods=["POST", "OPTIONS"])
def methodology():
    def gen():
        body          = request.json or {}
        topic         = body.get("topic", "").strip()
        research_type = body.get("research_type", "Qualitative").strip()
        collection    = body.get("collection", "Interviews").strip()

        if not topic:
            yield f"data: {json.dumps({'error': 'Topic is required'})}\n\n"
            return

        prompt = f"""Write a complete Chapter 3 Methodology section for:

Topic: {topic}
Research Type: {research_type}
Data Collection Method: {collection}

Include these sub-sections:
1. Research Philosophy / Paradigm
2. Research Design
3. Population and Sampling Strategy
4. Data Collection Instruments
5. Data Analysis Approach
6. Validity, Reliability and Ethical Considerations

Write in formal academic English. Minimum 500 words."""

        yield from stream_generator(SYSTEM_BASE, prompt, task="best_quality", max_tokens=1600)

    return sse_endpoint(gen)

# 6. Academic Editor
@app.route("/api/edit", methods=["POST", "OPTIONS"])
def edit():
    def gen():
        body   = request.json or {}
        text   = body.get("text", "").strip()
        action = body.get("action", "improve").strip()

        if not text:
            yield f"data: {json.dumps({'error': 'Text is required'})}\n\n"
            return

        action_prompts = {
            "improve":  "Rewrite the following text with a stronger academic tone, formal vocabulary, and improved sentence structure. Preserve all key ideas.",
            "shorten":  "Condense the following text to roughly half its length while keeping all critical points and maintaining academic style.",
            "expand":   "Expand the following text with additional explanation, academic context, and supporting detail. Keep a formal tone.",
            "simplify": "Rewrite the following text in clearer, simpler language while maintaining an appropriate academic register.",
        }

        instruction = action_prompts.get(action, action_prompts["improve"])
        prompt = f"""{instruction}

TEXT:
{text}

Return only the rewritten text — no preamble, no commentary."""

        yield from stream_generator(SYSTEM_BASE, prompt, task="best_quality", max_tokens=1500)

    return sse_endpoint(gen)

# ── Word (.docx) Export ────────────────────────────────────────────
@app.route("/api/export/docx", methods=["POST", "OPTIONS"])
def export_docx():
    if request.method == "OPTIONS":
        return "", 200

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import re

    body = request.json or {}
    sections = body.get("sections", {})

    SECTION_LABELS = {
        "topics":      "Dissertation Topics",
        "questions":   "Research Questions",
        "outline":     "Dissertation Outline",
        "writer":      "Written Section",
        "methodology": "Methodology",
        "editor":      "Edited Text",
    }

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("t3n28 Dissertation AI")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 0, 0)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Generated Research Content")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    for key, label in SECTION_LABELS.items():
        content = sections.get(key, "").strip()
        if not content:
            continue

        heading = doc.add_heading(label, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        heading.runs[0].font.size = Pt(14)

        for para_text in content.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            is_list = bool(re.match(r'^(\d+[\.\)]\s|[•\-]\s)', para_text))
            p = doc.add_paragraph(style='List Number' if is_list else 'Normal')
            clean = re.sub(r'^(\d+[\.\)]\s|[•\-]\s)', '', para_text) if is_list else para_text
            run = p.add_run(clean)
            run.font.size = Pt(11)
            if not is_list:
                p.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        buf.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': 'attachment; filename="t3n28-dissertation.docx"',
            'Access-Control-Allow-Origin': '*',
        }
    )

# ── Utility ────────────────────────────────────────────────────────
@app.route("/models")
def list_models():
    return jsonify({
        "working_models": WORKING_MODELS,
        "categories": MODEL_CATEGORIES,
        "stats": model_stats,
        "total_working": len(WORKING_MODELS)
    })

@app.route("/models/reset-stats", methods=["POST"])
def reset_stats():
    global model_stats
    model_stats = {m: {"successes": 0, "failures": 0, "avg_time": 0} for m in WORKING_MODELS}
    return jsonify({"success": True})

if __name__ == "__main__":
    app.run(debug=True, host="127.0.0.1", port=5000)
